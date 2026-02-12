"""
Service for managing session exports
"""
import logging
import json
import zipfile
import io
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional, List
from xml.etree import ElementTree as ET
from xml.dom import minidom

from config import (
    EXPORTS_DIR,
    ENABLE_AUTO_EXPORT
)
from my_app.utils.helpers import (
    generate_timestamp,
    calculate_node_execution_order
)
from my_app.services.comfyui_service import comfyui_service

logger = logging.getLogger(__name__)

# Import export libraries
try:
    import logging as weasy_logging
    from weasyprint import HTML, CSS
    # Suppress WeasyPrint CSS warnings
    weasy_logging.getLogger('weasyprint').setLevel(weasy_logging.ERROR)
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False
    logger.warning("WeasyPrint not available - PDF export disabled")

try:
    from docx import Document
    from docx.shared import Inches
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX export disabled")


class ExportManager:
    """Manager for session exports"""
    
    def __init__(self):
        self.exports_dir = EXPORTS_DIR
        self.exports_dir.mkdir(exist_ok=True)
        
        # Create subdirectories for different export formats
        self.html_dir = self.exports_dir / "html"
        self.html_dir.mkdir(exist_ok=True)
        
        self.pdf_dir = self.exports_dir / "pdf"
        self.pdf_dir.mkdir(exist_ok=True)
        
        self.xml_dir = self.exports_dir / "xml"
        self.xml_dir.mkdir(exist_ok=True)
        
        self.docx_dir = self.exports_dir / "docx"
        self.docx_dir.mkdir(exist_ok=True)
    
    def generate_export_html(self, session_data: Dict[str, Any], media_files: List[str]) -> str:
        """Generate HTML file for a session export"""
        user_id = session_data.get('user_id', 'DOE_J')
        timestamp = session_data.get('timestamp', generate_timestamp())
        session_id = session_data.get('session_id', 'unknown')
        workflow_name = session_data.get('workflow_name', 'Unknown Workflow')
        prompt = session_data.get('prompt', '')
        translated_prompt = session_data.get('translated_prompt', '')
        used_seed = session_data.get('used_seed', 'N/A')
        safety_level = session_data.get('safety_level', 'research')
        outputs = session_data.get('outputs', [])
        
        html_content = f"""<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AI4ArtsEd Export - {user_id}_{timestamp}_{session_id}</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif; 
               background-color: #f0f2f5; color: #1c1e21; padding: 20px; line-height: 1.6; }}
        .container {{ max-width: 800px; margin: 0 auto; background-color: #ffffff;
                     border-radius: 8px; padding: 24px; border: 1px solid #ddd; }}
        h1 {{ color: #333; border-bottom: 2px solid #007bff; padding-bottom: 12px; }}
        .metadata {{ background-color: #f8f9fa; padding: 15px; border-radius: 6px; margin-bottom: 20px; border: 1px solid #e0e0e0; }}
        .metadata-item {{ margin-bottom: 8px; }}
        .metadata-label {{ font-weight: bold; color: #495057; }}
        .prompt-section {{ background-color: #e7f3ff; padding: 15px; border-radius: 6px; margin-bottom: 20px; border: 1px solid #b8daff; }}
        .output-item {{ margin-bottom: 30px; padding: 15px; border: 1px solid #ddd; 
                       border-radius: 6px; background-color: #f9f9f9; }}
        .output-header {{ font-weight: bold; color: #555; margin-bottom: 10px; font-size: 1.1em; }}
        .output-content img {{ max-width: 100%; height: auto; border-radius: 8px; border: 1px solid #ddd; }}
        .output-content audio {{ width: 100%; margin-top: 10px; }}
        .text-content {{ white-space: pre-wrap; word-wrap: break-word; font-family: monospace; 
                        background: white; padding: 12px; border-radius: 4px; border: 1px solid #ddd; }}
        .footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #eee; 
                  text-align: center; color: #666; font-size: 0.9em; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>AI4ArtsEd Workflow Export</h1>
        
        <div class="metadata">
            <div class="metadata-item"><span class="metadata-label">User ID:</span> {user_id}</div>
            <div class="metadata-item"><span class="metadata-label">Session ID:</span> {session_id}</div>
            <div class="metadata-item"><span class="metadata-label">Timestamp:</span> {timestamp}</div>
            <div class="metadata-item"><span class="metadata-label">Workflow:</span> {workflow_name}</div>
            <div class="metadata-item"><span class="metadata-label">Seed:</span> {used_seed}</div>
            <div class="metadata-item"><span class="metadata-label">Filter-Status:</span> {safety_level}</div>
            <div class="metadata-item"><span class="metadata-label">Export Date:</span> {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</div>
        </div>
        
        <div class="prompt-section">
            <h3>Original Prompt</h3>
            <div class="text-content">{prompt}</div>
        </div>
        
        {f'''<div class="prompt-section" style="background-color: #f0f8ff;">
            <h3>Übersetzter Prompt</h3>
            <div class="text-content">{translated_prompt}</div>
        </div>''' if translated_prompt and translated_prompt != prompt else ''}
        
        <h3>Workflow Outputs (in Processing Order)</h3>
"""
        
        for i, output in enumerate(outputs, 1):
            html_content += f"""
        <div class="output-item">
            <div class="output-header">{i}. {output['title']} ({output['type'].upper()})</div>
            <div class="output-content">
"""
            
            if output['type'] == 'image':
                filename = output.get('filename', f"image_{i:03d}.png")
                html_content += f'<img src="media/{filename}" alt="Generated Image {i}">'
            elif output['type'] == 'text':
                html_content += f'<div class="text-content">{output["content"]}</div>'
            elif output['type'] == 'audio':
                filename = output.get('filename', f"audio_{i:03d}.wav")
                html_content += f'<audio controls><source src="media/{filename}" type="audio/wav">Your browser does not support the audio element.</audio>'
            
            html_content += """
            </div>
        </div>
"""
        
        html_content += f"""
        <div class="footer">
            <p>Generated by AI4ArtsEd - Artificial Intelligence for Arts Education</p>
            <p>Export created: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}</p>
        </div>
    </div>
</body>
</html>"""
        
        return html_content
    
    def process_outputs(self, outputs: Dict[str, Any], workflow_def: Dict[str, Any], 
                       media_dir: Path, session_prefix: str = None, 
                       workflow_name: str = None) -> tuple:
        """
        Process workflow outputs and download media files
        
        Args:
            outputs: ComfyUI outputs
            workflow_def: Workflow definition
            media_dir: Directory to save media files
            session_prefix: Prefix for media filenames (e.g., "session_DOE_J_250714202907_206a5e28")
            workflow_name: Workflow name to include in filenames
            
        Returns:
            Tuple of (processed_outputs, media_files)
        """
        processed_outputs = []
        media_files = []
        image_counter = 0
        audio_counter = 0
        
        # DEBUG: Log outputs
        logger.info(f"[EXPORT DEBUG] process_outputs called with {len(outputs)} outputs")
        logger.info(f"[EXPORT DEBUG] Media dir: {media_dir}")
        
        # Clean workflow name for filename
        workflow_clean = workflow_name.replace('.json', '') if workflow_name else 'unknown'
        
        for node_id, output in outputs.items():
            node_title = workflow_def.get(node_id, {}).get("_meta", {}).get("title", f"Node {node_id}")
            execution_order = calculate_node_execution_order(node_id, workflow_def)
            
            if output.get("text"):
                processed_outputs.append({
                    "title": node_title,
                    "type": "text",
                    "content": "\n".join(output["text"]),
                    "execution_order": execution_order
                })
            
            if output.get("images"):
                logger.info(f"[EXPORT DEBUG] Node {node_id} has {len(output['images'])} images")
                for idx, img in enumerate(output["images"]):
                    image_counter += 1
                    if session_prefix and workflow_name:
                        filename = f"{session_prefix}_{workflow_clean}_image_{image_counter:03d}.png"
                    else:
                        filename = f"image_{image_counter:03d}.png"
                    
                    img_url = f"view?filename={img['filename']}&subfolder={img['subfolder']}&type={img['type']}"
                    logger.info(f"[EXPORT DEBUG] Downloading image from: {img_url} to {media_dir / filename}")
                    
                    # Download image
                    if comfyui_service.download_media(img_url, media_dir / filename):
                        logger.info(f"[EXPORT DEBUG] Successfully downloaded {filename}")
                        processed_outputs.append({
                            "title": node_title,
                            "type": "image",
                            "filename": filename,
                            "execution_order": execution_order
                        })
                        media_files.append(filename)
                    else:
                        logger.error(f"[EXPORT DEBUG] Failed to download {filename}")
            
            if output.get("audio"):
                for idx, aud in enumerate(output["audio"]):
                    audio_counter += 1
                    if session_prefix and workflow_name:
                        filename = f"{session_prefix}_{workflow_clean}_audio_{audio_counter:03d}.wav"
                    else:
                        filename = f"audio_{audio_counter:03d}.wav"
                    
                    aud_url = f"view?filename={aud['filename']}&subfolder={aud['subfolder']}&type={aud['type']}"
                    
                    # Download audio
                    if comfyui_service.download_media(aud_url, media_dir / filename):
                        processed_outputs.append({
                            "title": node_title,
                            "type": "audio",
                            "filename": filename,
                            "execution_order": execution_order
                        })
                        media_files.append(filename)
        
        # Sort outputs by execution order
        processed_outputs.sort(key=lambda x: x["execution_order"])
        
        logger.info(f"[EXPORT DEBUG] Processed {len(processed_outputs)} outputs, {len(media_files)} media files")
        
        return processed_outputs, media_files
    
    def auto_export_session(self, prompt_id: str, workflow_name: str, prompt_text: str,
                          translated_prompt: str = None, used_seed: int = None, 
                          safety_level: str = 'research') -> bool:
        """
        Automatically export a session after successful completion
        
        Returns:
            True if export successful, False otherwise
        """
        if not ENABLE_AUTO_EXPORT:
            logger.info(f"Auto-export disabled for session {prompt_id}")
            return False
        
        try:
            # Get session data from ComfyUI - EXACT COPY FROM MANUAL EXPORT
            result = comfyui_service.get_workflow_outputs(prompt_id)
            if not result:
                return False
            
            outputs = result["outputs"]
            workflow_def = result["workflow_def"]
            
            # Generate session identifiers
            timestamp = generate_timestamp()
            session_id = prompt_id[:8]
            user_id = "DOE_J"  # Default user for auto-export
            
            # Create export directory structure in html subdirectory
            session_dir_name = f"session_{user_id}_{timestamp}_{session_id}"
            session_dir = self.html_dir / session_dir_name
            session_dir.mkdir(exist_ok=True)
            
            media_dir = session_dir / "media"
            media_dir.mkdir(exist_ok=True)
            
            # Process outputs and collect media - USING CODE FROM create_download_zip
            processed_outputs = []
            media_files = []
            media_counter = 1
            
            # Clean workflow name for filename - remove path separators
            workflow_clean = workflow_name.replace('.json', '').replace('/', '_').replace('\\', '_')
            
            for node_id, output in outputs.items():
                node_title = workflow_def.get(node_id, {}).get("_meta", {}).get("title", f"Node {node_id}")
                execution_order = calculate_node_execution_order(node_id, workflow_def)
                
                if output.get("text"):
                    processed_outputs.append({
                        "title": node_title,
                        "type": "text",
                        "content": "\n".join(output["text"]),
                        "execution_order": execution_order
                    })
                
                if output.get("images"):
                    for idx, img in enumerate(output["images"]):
                        filename = f"{session_dir_name}_{workflow_clean}_image_{media_counter:03d}.png"
                        img_url = f"view?filename={img['filename']}&subfolder={img['subfolder']}&type={img['type']}"
                        
                        # Download directly using proxy_request like create_download_zip
                        try:
                            response = comfyui_service.proxy_request(img_url)
                            if response.status_code == 200:
                                # Save to media directory
                                with open(media_dir / filename, 'wb') as f:
                                    f.write(response.content)
                                
                                processed_outputs.append({
                                    "title": node_title,
                                    "type": "image",
                                    "filename": filename,
                                    "execution_order": execution_order
                                })
                                media_files.append(filename)
                                media_counter += 1
                                logger.info(f"Auto-export: Downloaded image {filename}")
                        except Exception as e:
                            logger.error(f"Failed to download image: {e}")
                
                if output.get("audio"):
                    for idx, aud in enumerate(output["audio"]):
                        filename = f"{session_dir_name}_{workflow_clean}_audio_{media_counter:03d}.wav"
                        aud_url = f"view?filename={aud['filename']}&subfolder={aud['subfolder']}&type={aud['type']}"
                        
                        # Download directly using proxy_request like create_download_zip
                        try:
                            response = comfyui_service.proxy_request(aud_url)
                            if response.status_code == 200:
                                # Save to media directory
                                with open(media_dir / filename, 'wb') as f:
                                    f.write(response.content)
                                
                                processed_outputs.append({
                                    "title": node_title,
                                    "type": "audio",
                                    "filename": filename,
                                    "execution_order": execution_order
                                })
                                media_files.append(filename)
                                media_counter += 1
                                logger.info(f"Auto-export: Downloaded audio {filename}")
                        except Exception as e:
                            logger.error(f"Failed to download audio: {e}")
            
            # Sort outputs by execution order
            processed_outputs.sort(key=lambda x: x["execution_order"])
            logger.info(f"Auto-export: Processed {len(processed_outputs)} outputs, {len(media_files)} media files")
            
            # Create session data - INCLUDING AUTO-EXPORT SPECIFIC FIELDS
            session_data = {
                "user_id": user_id,
                "session_id": session_id,
                "timestamp": timestamp,
                "workflow_name": workflow_name,
                "prompt": prompt_text,
                "translated_prompt": translated_prompt or prompt_text,  # Auto-export specific
                "used_seed": used_seed,  # Auto-export specific
                "safety_level": safety_level,  # Auto-export specific
                "outputs": processed_outputs,
                "session_dir_name": session_dir_name
            }
            
            
            # Generate HTML file with workflow name
            html_filename = f"output_{user_id}_{timestamp}_{session_id}_{workflow_clean}.html"
            html_content = self.generate_export_html(session_data, media_files)
            
            with open(session_dir / html_filename, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Generate PDF export
            pdf_filename = f"output_{user_id}_{timestamp}_{session_id}_{workflow_clean}.pdf"
            self.generate_export_pdf(html_content, session_data, self.pdf_dir / pdf_filename)
            
            # Generate XML export
            xml_filename = f"export_{user_id}_{timestamp}_{session_id}_{workflow_clean}.xml"
            self.generate_export_xml(session_data, self.xml_dir / xml_filename, session_dir)
            
            # Generate DOCX export
            docx_filename = f"export_{user_id}_{timestamp}_{session_id}_{workflow_clean}.docx"
            self.generate_export_docx(session_data, self.docx_dir / docx_filename, session_dir)
            
            # Create metadata.json
            metadata = {
                "user_id": user_id,
                "session_id": session_id,
                "timestamp": timestamp,
                "workflow_name": workflow_name,
                "prompt": prompt_text,
                "translated_prompt": translated_prompt or prompt_text,  # Auto-export specific
                "used_seed": used_seed,  # Auto-export specific
                "safety_level": safety_level,  # Auto-export specific
                "export_date": datetime.now().isoformat(),
                "media_files": media_files,
                "output_count": len(processed_outputs)
            }
            
            with open(session_dir / "metadata.json", 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Successfully auto-exported session {session_id} to {session_dir_name}")
            
            # Update the sessions overview
            self._update_sessions_js()
            
            return True
            
        except Exception as e:
            logger.error(f"Auto-export failed: {e}")
            return False
    
    def export_session(self, prompt_id: str, user_id: str, workflow_name: str, 
                      prompt_text: str) -> Dict[str, Any]:
        """
        Manually export a session
        
        Returns:
            Export result dictionary
        """
        try:
            # Get session data from ComfyUI
            result = comfyui_service.get_workflow_outputs(prompt_id)
            if not result:
                return {"success": False, "error": "Session nicht in ComfyUI History gefunden."}
            
            outputs = result["outputs"]
            workflow_def = result["workflow_def"]
            
            # Generate session identifiers
            timestamp = generate_timestamp()
            session_id = prompt_id[:8]
            
            # Create export directory structure in html subdirectory
            session_dir_name = f"session_{user_id}_{timestamp}_{session_id}"
            session_dir = self.html_dir / session_dir_name
            session_dir.mkdir(exist_ok=True)
            
            media_dir = session_dir / "media"
            media_dir.mkdir(exist_ok=True)
            
            # Process outputs and collect media with new naming convention
            processed_outputs, media_files = self.process_outputs(
                outputs, workflow_def, media_dir,
                session_prefix=session_dir_name,
                workflow_name=workflow_name
            )
            
            # Create session data
            session_data = {
                "user_id": user_id,
                "session_id": session_id,
                "timestamp": timestamp,
                "workflow_name": workflow_name,
                "prompt": prompt_text,
                "outputs": processed_outputs,
                "session_dir_name": session_dir_name
            }
            
            # Clean workflow name for filename
            workflow_clean = workflow_name.replace('.json', '')
            
            # Generate HTML file with workflow name
            html_filename = f"output_{user_id}_{timestamp}_{session_id}_{workflow_clean}.html"
            html_content = self.generate_export_html(session_data, media_files)
            
            with open(session_dir / html_filename, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Generate PDF export
            pdf_filename = f"output_{user_id}_{timestamp}_{session_id}_{workflow_clean}.pdf"
            self.generate_export_pdf(html_content, session_data, self.pdf_dir / pdf_filename)
            
            # Generate XML export
            xml_filename = f"export_{user_id}_{timestamp}_{session_id}_{workflow_clean}.xml"
            self.generate_export_xml(session_data, self.xml_dir / xml_filename, session_dir)
            
            # Generate DOCX export
            docx_filename = f"export_{user_id}_{timestamp}_{session_id}_{workflow_clean}.docx"
            self.generate_export_docx(session_data, self.docx_dir / docx_filename, session_dir)
            
            # Create metadata.json
            metadata = {
                "user_id": user_id,
                "session_id": session_id,
                "timestamp": timestamp,
                "workflow_name": workflow_name,
                "prompt": prompt_text,
                "export_date": datetime.now().isoformat(),
                "media_files": media_files,
                "output_count": len(processed_outputs)
            }
            
            with open(session_dir / "metadata.json", 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False)
            
            logger.info(f"Successfully exported session {session_id} to {session_dir_name}")
            
            # Update the sessions overview
            self._update_sessions_js()
            
            return {
                "success": True,
                "export_path": session_dir_name,
                "html_file": html_filename,
                "media_count": len(media_files),
                "output_count": len(processed_outputs),
                "pdf_file": pdf_filename,
                "xml_file": xml_filename,
                "docx_file": docx_filename
            }
            
        except Exception as e:
            logger.error(f"Export failed: {e}")
            return {"success": False, "error": f"Export fehlgeschlagen: {str(e)}"}
    
    def create_download_zip(self, prompt_id: str, user_id: str, workflow_name: str, 
                           prompt_text: str) -> Optional[bytes]:
        """
        Create a ZIP file for download
        
        Returns:
            ZIP file bytes or None if failed
        """
        try:
            # Get session data from ComfyUI
            result = comfyui_service.get_workflow_outputs(prompt_id)
            if not result:
                return None
            
            outputs = result["outputs"]
            workflow_def = result["workflow_def"]
            
            # Generate session identifiers
            timestamp = generate_timestamp()
            session_id = prompt_id[:8]
            
            # Process outputs and collect media
            processed_outputs = []
            
            # Create ZIP file in memory
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                media_counter = 1
                
                for node_id, output in outputs.items():
                    node_title = workflow_def.get(node_id, {}).get("_meta", {}).get("title", f"Node {node_id}")
                    execution_order = calculate_node_execution_order(node_id, workflow_def)
                    
                    if output.get("text"):
                        processed_outputs.append({
                            "title": node_title,
                            "type": "text",
                            "content": "\n".join(output["text"]),
                            "execution_order": execution_order
                        })
                    
                    if output.get("images"):
                        for idx, img in enumerate(output["images"]):
                            filename = f"image_{media_counter:03d}.png"
                            img_url = f"view?filename={img['filename']}&subfolder={img['subfolder']}&type={img['type']}"
                            
                            # Download and add to ZIP
                            try:
                                response = comfyui_service.proxy_request(img_url)
                                if response.status_code == 200:
                                    zip_file.writestr(f"media/{filename}", response.content)
                                    processed_outputs.append({
                                        "title": node_title,
                                        "type": "image",
                                        "filename": filename,
                                        "execution_order": execution_order
                                    })
                                    media_counter += 1
                            except Exception as e:
                                logger.error(f"Failed to add image to ZIP: {e}")
                    
                    if output.get("audio"):
                        for idx, aud in enumerate(output["audio"]):
                            filename = f"audio_{media_counter:03d}.wav"
                            aud_url = f"view?filename={aud['filename']}&subfolder={aud['subfolder']}&type={aud['type']}"
                            
                            # Download and add to ZIP
                            try:
                                response = comfyui_service.proxy_request(aud_url)
                                if response.status_code == 200:
                                    zip_file.writestr(f"media/{filename}", response.content)
                                    processed_outputs.append({
                                        "title": node_title,
                                        "type": "audio",
                                        "filename": filename,
                                        "execution_order": execution_order
                                    })
                                    media_counter += 1
                            except Exception as e:
                                logger.error(f"Failed to add audio to ZIP: {e}")
                
                # Sort outputs by execution order
                processed_outputs.sort(key=lambda x: x["execution_order"])
                
                # Create session data
                session_data = {
                    "user_id": user_id,
                    "session_id": session_id,
                    "timestamp": timestamp,
                    "workflow_name": workflow_name,
                    "prompt": prompt_text,
                    "outputs": processed_outputs
                }
                
                # Generate HTML file and add to ZIP
                html_filename = f"output_{user_id}_{timestamp}_{session_id}.html"
                html_content = self.generate_export_html(session_data, [])
                zip_file.writestr(html_filename, html_content.encode('utf-8'))
                
                # Create metadata.json and add to ZIP
                metadata = {
                    "user_id": user_id,
                    "session_id": session_id,
                    "timestamp": timestamp,
                    "workflow_name": workflow_name,
                    "prompt": prompt_text,
                    "export_date": datetime.now().isoformat(),
                    "media_files": [out.get("filename") for out in processed_outputs if out.get("filename")],
                    "output_count": len(processed_outputs)
                }
                zip_file.writestr("metadata.json", json.dumps(metadata, indent=2, ensure_ascii=False).encode('utf-8'))
            
            zip_buffer.seek(0)
            logger.info(f"Successfully created ZIP download for session {session_id}")
            
            return zip_buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to create download ZIP: {e}")
            return None


    def _update_sessions_js(self):
        """
        Scans the export directory and updates a sessions.js file
        for the local overview page.
        """
        try:
            sessions = []
            # Look for sessions in both old location and new html subdirectory
            for location in [self.exports_dir, self.html_dir]:
                for item in location.iterdir():
                    if item.is_dir() and item.name.startswith('session_'):
                        metadata_path = item / 'metadata.json'
                        if metadata_path.exists():
                            with open(metadata_path, 'r', encoding='utf-8') as f:
                                data = json.load(f)
                        
                        user_id = data.get('user_id', 'N/A')
                        session_id = data.get('session_id', 'N/A')
                        timestamp = data.get('timestamp', 'N/A')
                        workflow_name = data.get('workflow_name', 'N/A')
                        workflow_clean = workflow_name.replace('.json', '').replace('/', '_').replace('\\', '_') if workflow_name != 'N/A' else 'N/A'
                        
                        # Look for actual HTML file in the directory
                        html_files = list(item.glob(f"output_{user_id}_{timestamp}_{session_id}_*.html"))
                        if html_files:
                            html_file = html_files[0].name
                        else:
                            # Fallback to old naming convention
                            html_file = f"output_{user_id}_{timestamp}_{session_id}.html"
                        
                        # Generate PDF, XML, DOCX filenames
                        pdf_file = f"output_{user_id}_{timestamp}_{session_id}_{workflow_clean}.pdf"
                        xml_file = f"export_{user_id}_{timestamp}_{session_id}_{workflow_clean}.xml"
                        docx_file = f"export_{user_id}_{timestamp}_{session_id}_{workflow_clean}.docx"

                        sessions.append({
                            'user_id': user_id,
                            'session_id': session_id,
                            'timestamp': timestamp,
                            'workflow_name': workflow_name,
                            'workflow_clean': workflow_clean,
                            'output_count': data.get('output_count', 0),
                            'media_count': len(data.get('media_files', [])),
                            'export_date': data.get('export_date'),
                            'folder_name': item.name,
                            'html_file': html_file,
                            'pdf_file': pdf_file,
                            'xml_file': xml_file,
                            'docx_file': docx_file
                        })
            
            # Write to sessions.js
            js_content = f"const allSessions = {json.dumps(sessions, indent=2, ensure_ascii=False)};"
            with open(self.exports_dir / "sessions.js", 'w', encoding='utf-8') as f:
                f.write(js_content)
            
            logger.info(f"Successfully updated sessions.js with {len(sessions)} sessions.")

        except Exception as e:
            logger.error(f"Failed to update sessions.js: {e}")
    
    def generate_export_pdf(self, html_content: str, session_data: Dict[str, Any], 
                           output_path: Path) -> bool:
        """
        Generate PDF from HTML content
        
        Returns:
            True if successful, False otherwise
        """
        if not WEASYPRINT_AVAILABLE:
            logger.warning("PDF export skipped - WeasyPrint not available")
            return False
        
        try:
            # Log the HTML content to debug potential box-shadow references
            logger.info(f"[PDF DEBUG] Generating PDF for session {session_data.get('session_id', 'unknown')}")
            
            # Create CSS for single-page PDF with explicit box-shadow removal
            css = CSS(string="""
                @page {
                    size: A4;
                    margin: 20mm;
                }
                body {
                    font-family: Arial, sans-serif !important;
                    background-color: white !important;
                    color: black !important;
                    line-height: 1.4 !important;
                }
                .container {
                    max-width: none !important;
                    background-color: white !important;
                    border: none !important;
                    padding: 10px !important;
                }
                * {
                    background: white !important;
                    border-radius: 0 !important;
                }
            """)
            
            # Generate PDF with metadata (no base_url to avoid CSS file conflicts)
            pdf = HTML(string=html_content).write_pdf(
                stylesheets=[css]
            )
            
            # Write PDF with metadata
            output_path.write_bytes(pdf)
            
            logger.info(f"PDF export successful: {output_path.name}")
            return True
            
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            return False
    
    def generate_export_xml(self, session_data: Dict[str, Any], output_path: Path, 
                           session_dir: Path) -> bool:
        """
        Generate XML export for QDA software
        
        Returns:
            True if successful, False otherwise
        """
        try:
            # Create root element
            root = ET.Element("export")
            session = ET.SubElement(root, "session")
            
            # Add metadata
            ET.SubElement(session, "id").text = session_data["session_dir_name"]
            ET.SubElement(session, "user").text = session_data["user_id"]
            ET.SubElement(session, "timestamp").text = session_data["timestamp"]
            ET.SubElement(session, "workflow").text = session_data["workflow_name"].replace('.json', '')
            ET.SubElement(session, "seed").text = str(session_data.get("used_seed", "N/A"))
            ET.SubElement(session, "safety_level").text = session_data.get("safety_level", "research")
            
            # Add prompts
            prompts = ET.SubElement(session, "prompts")
            ET.SubElement(prompts, "original").text = session_data["prompt"]
            ET.SubElement(prompts, "translated").text = session_data.get("translated_prompt", session_data["prompt"])
            
            # Add outputs
            outputs_elem = ET.SubElement(session, "outputs")
            for idx, output in enumerate(session_data["outputs"], 1):
                output_elem = ET.SubElement(outputs_elem, "output")
                output_elem.set("order", str(idx))
                output_elem.set("type", output["type"])
                output_elem.set("node_title", output["title"])
                
                if output["type"] == "text":
                    ET.SubElement(output_elem, "content").text = output["content"]
                elif output["type"] in ["image", "audio"]:
                    ET.SubElement(output_elem, "filename").text = output["filename"]
                    ET.SubElement(output_elem, "path").text = f"{session_data['session_dir_name']}/media/{output['filename']}"
            
            # Pretty print XML
            xml_str = minidom.parseString(ET.tostring(root)).toprettyxml(indent="  ")
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(xml_str)
            
            logger.info(f"XML export successful: {output_path.name}")
            return True
            
        except Exception as e:
            logger.error(f"XML export failed: {e}")
            return False
    
    def generate_export_docx(self, session_data: Dict[str, Any], output_path: Path, 
                            session_dir: Path) -> bool:
        """
        Generate DOCX export for QDA software
        
        Returns:
            True if successful, False otherwise
        """
        if not PYTHON_DOCX_AVAILABLE:
            logger.warning("DOCX export skipped - python-docx not available")
            return False
        
        try:
            doc = Document()
            
            # Add title
            doc.add_heading('AI4ArtsEd Workflow Export', 0)
            
            # Add metadata table
            doc.add_heading('Session Information', level=1)
            table = doc.add_table(rows=7, cols=2)
            table.style = 'Light List Accent 1'
            
            metadata_items = [
                ('User ID', session_data["user_id"]),
                ('Session ID', session_data["session_id"]),
                ('Timestamp', session_data["timestamp"]),
                ('Workflow', session_data["workflow_name"]),
                ('Seed', str(session_data.get("used_seed", "N/A"))),
                ('Filter-Status', session_data.get("safety_level", "research")),
                ('Export Date', datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            ]
            
            for i, (label, value) in enumerate(metadata_items):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = value
            
            # Add prompts
            doc.add_heading('Prompts', level=1)
            doc.add_heading('Original Prompt', level=2)
            doc.add_paragraph(session_data["prompt"])
            
            if session_data.get("translated_prompt") and session_data["translated_prompt"] != session_data["prompt"]:
                doc.add_heading('Translated Prompt', level=2)
                doc.add_paragraph(session_data["translated_prompt"])
            
            # Add outputs
            doc.add_heading('Workflow Outputs (in Processing Order)', level=1)
            
            for idx, output in enumerate(session_data["outputs"], 1):
                doc.add_heading(f'{idx}. {output["title"]} ({output["type"].upper()})', level=2)
                
                if output["type"] == "text":
                    doc.add_paragraph(output["content"])
                elif output["type"] == "image":
                    try:
                        img_path = session_dir / "media" / output["filename"]
                        if img_path.exists():
                            doc.add_picture(str(img_path), width=Inches(6))
                            doc.add_paragraph(f'Image: {output["filename"]}', style='Caption')
                    except Exception as e:
                        doc.add_paragraph(f'[Image: {output["filename"]} - Could not embed]')
                        logger.warning(f"Could not embed image in DOCX: {e}")
                elif output["type"] == "audio":
                    doc.add_paragraph(f'Audio file: {output["filename"]}')
            
            # Add footer
            doc.add_page_break()
            footer_para = doc.add_paragraph()
            footer_para.alignment = 1  # Center alignment
            footer_para.add_run('Generated by AI4ArtsEd - Artificial Intelligence for Arts Education\n').italic = True
            footer_para.add_run(f'Export created: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}').italic = True
            
            # Save document
            doc.save(str(output_path))
            
            logger.info(f"DOCX export successful: {output_path.name}")
            return True
            
        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            return False

# Create a singleton instance
export_manager = ExportManager()
